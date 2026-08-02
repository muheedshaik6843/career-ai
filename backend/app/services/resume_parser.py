"""
Resume Parser Service
Supports PDF (via PyMuPDF) and DOCX (via python-docx)
Falls back to basic text parsing when AI libraries not available
"""
import re
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text("text") + "\n"
        doc.close()
        return text.strip()
    except ImportError:
        logger.warning("PyMuPDF not installed — using fallback PDF parser.")
        return _fallback_pdf_text(file_bytes)
    except Exception as e:
        logger.error(f"PDF parsing failed: {e}")
        raise ValueError(f"Could not parse PDF: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        logger.warning("python-docx not installed — cannot parse DOCX.")
        raise ValueError("DOCX parsing requires python-docx library.")
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        raise ValueError(f"Could not parse DOCX: {e}")


def _fallback_pdf_text(file_bytes: bytes) -> str:
    """Minimal PDF text extraction without external libraries."""
    try:
        text = file_bytes.decode("latin-1", errors="ignore")
        # Extract readable ASCII strings
        readable = re.findall(r'[\x20-\x7E]{4,}', text)
        return "\n".join(readable[:500])
    except Exception:
        return ""


# ─── NLP Parsing ─────────────────────────────────────────────────────────────

def extract_email(text: str) -> Optional[str]:
    pattern = r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(pattern, text)
    return matches[0] if matches else None


def extract_phone(text: str) -> Optional[str]:
    pattern = r'(\+?\d{1,3}[-.\s]?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4})'
    matches = re.findall(pattern, text)
    return matches[0].strip() if matches else None


def extract_name(text: str) -> Optional[str]:
    """Heuristic: first non-empty line that looks like a person's name."""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines[:5]:
        words = line.split()
        if 2 <= len(words) <= 5 and all(w[0].isupper() for w in words if w.isalpha()):
            return line
    return lines[0] if lines else None


# Common tech skills vocabulary
SKILLS_VOCABULARY = {
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust", "ruby",
    "php", "swift", "kotlin", "scala", "r", "matlab",
    "react", "next.js", "nextjs", "vue", "angular", "svelte", "nuxt",
    "node.js", "nodejs", "express", "fastapi", "django", "flask", "spring", "rails",
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "sqlite",
    "aws", "gcp", "azure", "docker", "kubernetes", "terraform", "ansible",
    "git", "github", "gitlab", "ci/cd", "github actions", "jenkins",
    "machine learning", "deep learning", "nlp", "tensorflow", "pytorch",
    "scikit-learn", "pandas", "numpy", "spark", "kafka",
    "graphql", "rest", "grpc", "websocket", "oauth", "jwt",
    "linux", "bash", "shell", "sql", "html", "css", "tailwindcss",
    "figma", "photoshop", "jira", "agile", "scrum", "microservices",
}


def extract_skills(text: str) -> list:
    """Extract tech skills from resume text using vocabulary matching."""
    text_lower = text.lower()
    found_skills = set()

    # Multi-word skills first
    for skill in SKILLS_VOCABULARY:
        if skill in text_lower:
            found_skills.add(skill.title() if len(skill) > 3 else skill.upper())

    # Also look for comma/bullet separated lists after "Skills:" heading
    skill_section_pattern = r'(?:skills?|technologies?|tech stack)[:\s]+([^\n]+(?:\n[^\n]+){0,5})'
    matches = re.findall(skill_section_pattern, text_lower, re.IGNORECASE)
    for match in matches:
        tokens = re.split(r'[,|•·\-\t]+', match)
        for token in tokens:
            t = token.strip()
            if 2 <= len(t) <= 40 and not t.isdigit():
                found_skills.add(t.title())

    return sorted(found_skills)[:50]


def extract_education(text: str) -> list:
    """Extract education sections."""
    education = []
    degree_patterns = [
        r'(B\.?S\.?|B\.?E\.?|B\.?Tech\.?|M\.?S\.?|M\.?E\.?|M\.?Tech\.?|MBA|Ph\.?D\.?|Bachelor|Master|Doctor)[^\n]*',
    ]
    university_keywords = ["university", "college", "institute", "school", "academy"]

    lines = text.split('\n')
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if any(kw in line_lower for kw in university_keywords):
            entry = {
                "institution": line.strip(),
                "degree": None,
                "field": None,
                "year": None,
            }
            # Look in surrounding lines for degree info
            context = '\n'.join(lines[max(0, i-2):i+3])
            for pattern in degree_patterns:
                m = re.search(pattern, context, re.IGNORECASE)
                if m:
                    entry["degree"] = m.group(0).strip()
                    break
            # Year
            year_match = re.search(r'\b(19|20)\d{2}\b', context)
            if year_match:
                entry["year"] = year_match.group(0)
            education.append(entry)

    return education[:5]


def extract_experience(text: str) -> list:
    """Extract work experience sections."""
    experience = []
    section_headers = ['experience', 'work history', 'employment', 'career', 'positions']

    in_section = False
    current_entry = None
    lines = text.split('\n')

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        line_lower = line_stripped.lower()
        if any(h in line_lower for h in section_headers) and len(line_stripped) < 40:
            in_section = True
            continue

        if in_section:
            # Detect new job entry (capitalized title-like line)
            if re.match(r'^[A-Z][a-zA-Z\s]{5,50}$', line_stripped):
                if current_entry:
                    experience.append(current_entry)
                current_entry = {
                    "title": line_stripped,
                    "company": None,
                    "duration": None,
                    "description": [],
                }
            elif current_entry:
                # Date range
                date_match = re.search(r'(\d{4}|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec).+(\d{4}|Present)', line_stripped, re.IGNORECASE)
                if date_match and not current_entry["duration"]:
                    current_entry["duration"] = line_stripped
                # Company name (short line after title)
                elif not current_entry["company"] and len(line_stripped) < 80:
                    current_entry["company"] = line_stripped
                else:
                    current_entry["description"].append(line_stripped)

    if current_entry:
        experience.append(current_entry)

    return experience[:10]


def extract_certifications(text: str) -> list:
    """Extract certifications."""
    certs = []
    cert_keywords = ['certified', 'certification', 'certificate', 'aws certified', 'google cloud', 'cpa', 'cfa', 'pmp', 'comptia']
    lines = text.split('\n')
    for line in lines:
        if any(kw in line.lower() for kw in cert_keywords):
            cert = line.strip()
            if cert and len(cert) < 200:
                certs.append(cert)
    return list(set(certs))[:10]


def extract_summary(text: str) -> Optional[str]:
    """Extract professional summary."""
    patterns = [
        r'(?:summary|profile|objective|about me)[:\s]+([^\n]+(?:\n[^\n]+){0,4})',
        r'^([A-Z][^.!?]*[.!?])\s*(?:[A-Z][^.!?]*[.!?]){1,3}'
    ]
    for pattern in patterns:
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            return m.group(1).strip()[:600]
    return None


def parse_resume(raw_text: str) -> dict:
    """Master parser — extracts all fields from raw text."""
    return {
        "full_name": extract_name(raw_text),
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "location": None,
        "summary": extract_summary(raw_text),
        "skills": extract_skills(raw_text),
        "education": extract_education(raw_text),
        "experience": extract_experience(raw_text),
        "projects": [],
        "certifications": extract_certifications(raw_text),
        "languages": [],
    }
